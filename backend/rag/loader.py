from __future__ import annotations

import csv
import json
import os
import re
from pathlib import Path
from typing import Any

from rag.records import RawRecord, serialize_records

try:
    import pypdf
except ImportError:
    pypdf = None

try:
    import docx
except ImportError:
    docx = None

try:
    import openpyxl
except ImportError:
    openpyxl = None


class Document:
    """Backward-compatible document wrapper.

    `content` remains a string so existing upload routes keep working. Structured
    formats are encoded as internal RAG records that `chunk_document()` understands.
    """

    def __init__(self, content: str, metadata: dict | None = None):
        self.content = content
        self.metadata = metadata or {}


ERROR_PREFIXES = (
    "[PDF Load Error:",
    "[PDF Read Error:",
    "[DOCX Load Error:",
    "[JSON Load Error:",
    "[CSV Load Error:",
    "[Excel Load Error:",
    "[Unsupported file format:",
)

SENSITIVE_FIELD_RE = re.compile(
    r"(?:password|passwd|secret|api[_ -]?key|access[_ -]?token|refresh[_ -]?token|"
    r"private[_ -]?key|otp|cvv|card[_ -]?number|auth[_ -]?token)",
    re.I,
)

IDENTITY_KEYS = (
    "title", "name", "question", "label", "id", "record_id", "sku", "code", "email"
)
CONTENT_KEYS = (
    "content", "text", "answer", "description", "details", "body", "summary", "value"
)
STRUCTURAL_KEYS = {
    "sample_questions", "keywords", "source_urls", "verification_status", "answer_mode",
    "risk_flags", "last_verified", "volatile", "source_type", "embedding", "metadata",
    "category", "topic", "company",
}


def _clean(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, bool):
        return "true" if value else "false"
    if isinstance(value, (dict, list)):
        return json.dumps(value, ensure_ascii=False, sort_keys=True)
    text = str(value).replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _safe_items(mapping: dict[str, Any]) -> list[tuple[str, Any]]:
    return [
        (str(key), value)
        for key, value in mapping.items()
        if not SENSITIVE_FIELD_RE.search(str(key))
        and value not in (None, "", [], {})
    ]


def _record_title(mapping: dict[str, Any], fallback: str) -> str:
    for key in IDENTITY_KEYS:
        value = mapping.get(key)
        if value not in (None, "", [], {}) and not isinstance(value, (dict, list)):
            return _clean(value)[:180]
    return fallback


def _mapping_to_text(mapping: dict[str, Any]) -> str:
    # Prefer an explicit natural-language content field, while retaining other useful fields.
    lines: list[str] = []
    used: set[str] = set()
    for key in CONTENT_KEYS:
        value = mapping.get(key)
        if value not in (None, "", [], {}):
            if isinstance(value, str):
                lines.append(_clean(value))
            else:
                lines.append(f"{key.replace('_', ' ').title()}: {_clean(value)}")
            used.add(key)
            break

    for key, value in _safe_items(mapping):
        if key in used or key in STRUCTURAL_KEYS or key in IDENTITY_KEYS:
            continue
        label = key.replace("_", " ").strip().title()
        if isinstance(value, dict):
            nested = "; ".join(
                f"{str(k).replace('_', ' ').title()}: {_clean(v)}"
                for k, v in _safe_items(value)
            )
            if nested:
                lines.append(f"{label}: {nested}")
        elif isinstance(value, list):
            if all(not isinstance(item, (dict, list)) for item in value):
                rendered = ", ".join(_clean(item) for item in value if _clean(item))
                if rendered:
                    lines.append(f"{label}: {rendered}")
            else:
                lines.append(f"{label}: {_clean(value)}")
        else:
            lines.append(f"{label}: {_clean(value)}")

    # Stable de-duplication.
    result: list[str] = []
    seen: set[str] = set()
    for line in lines:
        normalized = line.lower().strip()
        if normalized and normalized not in seen:
            seen.add(normalized)
            result.append(line.strip())
    return "\n".join(result)


def load_txt(file_path: str) -> str:
    with open(file_path, "r", encoding="utf-8", errors="ignore") as file:
        return file.read()


def load_pdf(file_path: str) -> str:
    if pypdf is None:
        return ""
    records: list[RawRecord] = []
    try:
        reader = pypdf.PdfReader(file_path)
        for page_index, page in enumerate(reader.pages, start=1):
            page_text = _clean(page.extract_text() or "")
            if not page_text:
                continue
            records.append(
                RawRecord(
                    title=f"{Path(file_path).stem} — Page {page_index}",
                    text=page_text,
                    source_type="pdf",
                    source_locator={"page_number": page_index},
                    record_id=f"page-{page_index}",
                )
            )
    except Exception:
        return ""
    return serialize_records(records)


def load_docx(file_path: str) -> str:
    if docx is None or Path(file_path).suffix.lower() == ".doc":
        return ""

    records: list[RawRecord] = []
    try:
        document = docx.Document(file_path)
        heading_path: list[str] = []
        buffer: list[str] = []
        current_title = Path(file_path).stem
        section_index = 0

        def flush() -> None:
            nonlocal buffer, section_index
            text = _clean("\n".join(buffer))
            if not text:
                buffer = []
                return
            section_index += 1
            records.append(
                RawRecord(
                    title=current_title,
                    text=text,
                    source_type="docx",
                    source_locator={
                        "section_index": section_index,
                        "heading_path": list(heading_path),
                    },
                    record_id=f"section-{section_index}",
                )
            )
            buffer = []

        for paragraph in document.paragraphs:
            text = _clean(paragraph.text)
            if not text:
                continue
            style_name = (getattr(paragraph.style, "name", "") or "").lower()
            match = re.match(r"heading\s+(\d+)", style_name)
            if match:
                flush()
                level = max(1, int(match.group(1)))
                heading_path[:] = heading_path[: level - 1]
                heading_path.append(text)
                current_title = " > ".join(heading_path)
            else:
                buffer.append(text)
        flush()

        for table_index, table in enumerate(document.tables, start=1):
            rows = [[_clean(cell.text) for cell in row.cells] for row in table.rows]
            rows = [row for row in rows if any(row)]
            if not rows:
                continue
            headers = [cell or f"Column {index + 1}" for index, cell in enumerate(rows[0])]
            for row_index, row in enumerate(rows[1:], start=2):
                fields = [
                    f"{headers[index]}: {value}"
                    for index, value in enumerate(row)
                    if value and index < len(headers)
                ]
                if fields:
                    records.append(
                        RawRecord(
                            title=f"{Path(file_path).stem} — Table {table_index}, Row {row_index}",
                            text="\n".join(fields),
                            source_type="docx",
                            source_locator={"table_index": table_index, "row_index": row_index},
                            record_id=f"table-{table_index}-row-{row_index}",
                        )
                    )
    except Exception:
        return ""
    return serialize_records(records)


def _json_records(value: Any, *, path: str, source_name: str) -> list[RawRecord]:
    records: list[RawRecord] = []

    if isinstance(value, list):
        for index, item in enumerate(value):
            item_path = f"{path}[{index}]" if path else f"[{index}]"
            if isinstance(item, dict):
                title = _record_title(item, f"{source_name} item {index + 1}")
                text = _mapping_to_text(item)
                if text:
                    records.append(
                        RawRecord(
                            title=title,
                            text=text,
                            source_type="json",
                            source_locator={"json_path": item_path, "json_index": index},
                            record_id=_clean(item.get("id") or item.get("record_id") or item_path),
                            metadata={
                                "sample_questions": item.get("sample_questions", []),
                                "provided_keywords": item.get("keywords", []),
                                "verification_status": item.get("verification_status"),
                                "answer_mode": item.get("answer_mode"),
                                "risk_flags": item.get("risk_flags", []),
                                "source_urls": item.get("source_urls", []),
                            },
                        )
                    )
            elif isinstance(item, list):
                records.extend(_json_records(item, path=item_path, source_name=source_name))
            else:
                text = _clean(item)
                if text:
                    records.append(
                        RawRecord(
                            title=f"{source_name} item {index + 1}",
                            text=text,
                            source_type="json",
                            source_locator={"json_path": item_path, "json_index": index},
                            record_id=item_path,
                        )
                    )
        return records

    if isinstance(value, dict):
        # Common KB layout: metadata + entries. Metadata is attached, entries become records.
        if isinstance(value.get("entries"), list):
            common_metadata = dict(value.get("metadata") or {})
            child_records = _json_records(value["entries"], path="entries", source_name=source_name)
            for record in child_records:
                record.metadata["document_metadata"] = common_metadata
            return child_records

        # Dictionaries whose values are mostly records are expanded; otherwise one coherent record.
        complex_values = [item for _, item in _safe_items(value) if isinstance(item, (dict, list))]
        scalar_values = [item for _, item in _safe_items(value) if not isinstance(item, (dict, list))]
        if complex_values and not scalar_values:
            for key, item in _safe_items(value):
                records.extend(
                    _json_records(
                        item,
                        path=f"{path}.{key}" if path else key,
                        source_name=f"{source_name} — {key}",
                    )
                )
            return records

        title = _record_title(value, source_name)
        text = _mapping_to_text(value)
        if text:
            records.append(
                RawRecord(
                    title=title,
                    text=text,
                    source_type="json",
                    source_locator={"json_path": path or "$"},
                    record_id=_clean(value.get("id") or value.get("record_id") or path or "$"),
                    metadata={
                        "sample_questions": value.get("sample_questions", []),
                        "provided_keywords": value.get("keywords", []),
                        "verification_status": value.get("verification_status"),
                        "answer_mode": value.get("answer_mode"),
                        "risk_flags": value.get("risk_flags", []),
                        "source_urls": value.get("source_urls", []),
                    },
                )
            )
        return records

    text = _clean(value)
    if text:
        records.append(
            RawRecord(
                title=source_name,
                text=text,
                source_type="json",
                source_locator={"json_path": path or "$"},
                record_id=path or "$",
            )
        )
    return records


def load_json(file_path: str) -> str:
    try:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as file:
            data = json.load(file)
    except Exception:
        return ""
    return serialize_records(
        _json_records(data, path="", source_name=Path(file_path).stem)
    )


def load_csv(file_path: str) -> str:
    records: list[RawRecord] = []
    try:
        with open(file_path, "r", encoding="utf-8-sig", errors="ignore", newline="") as file:
            sample = file.read(4096)
            file.seek(0)
            try:
                dialect = csv.Sniffer().sniff(sample)
            except Exception:
                dialect = csv.excel
            reader = csv.DictReader(file, dialect=dialect)
            headers = [
                (header.strip() if header and header.strip() else f"Column {index + 1}")
                for index, header in enumerate(reader.fieldnames or [])
            ]
            for row_index, raw_row in enumerate(reader, start=2):
                row = {
                    headers[index] if index < len(headers) else str(key): value
                    for index, (key, value) in enumerate(raw_row.items())
                    if key is not None and not SENSITIVE_FIELD_RE.search(str(key))
                }
                fields = [
                    f"{key}: {_clean(value)}"
                    for key, value in row.items()
                    if _clean(value)
                ]
                if not fields:
                    continue
                title = _record_title(row, f"{Path(file_path).stem} — Row {row_index}")
                records.append(
                    RawRecord(
                        title=title,
                        text="\n".join(fields),
                        source_type="csv",
                        source_locator={"row_index": row_index},
                        record_id=_clean(row.get("id") or row.get("record_id") or row_index),
                    )
                )
    except Exception:
        return ""
    return serialize_records(records)


def load_xlsx(file_path: str) -> str:
    if openpyxl is None or Path(file_path).suffix.lower() == ".xls":
        return ""

    records: list[RawRecord] = []
    try:
        workbook = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
        for sheet_name in workbook.sheetnames:
            worksheet = workbook[sheet_name]
            iterator = worksheet.iter_rows(values_only=True)
            try:
                first_row = next(iterator)
            except StopIteration:
                continue
            headers = [
                _clean(cell) or f"Column {index + 1}"
                for index, cell in enumerate(first_row)
            ]
            for row_index, row in enumerate(iterator, start=2):
                mapping = {
                    headers[index] if index < len(headers) else f"Column {index + 1}": cell
                    for index, cell in enumerate(row)
                    if cell not in (None, "")
                    and not SENSITIVE_FIELD_RE.search(
                        headers[index] if index < len(headers) else f"Column {index + 1}"
                    )
                }
                fields = [f"{key}: {_clean(value)}" for key, value in mapping.items() if _clean(value)]
                if not fields:
                    continue
                title = _record_title(mapping, f"{sheet_name} — Row {row_index}")
                records.append(
                    RawRecord(
                        title=title,
                        text="\n".join(fields),
                        source_type="xlsx",
                        source_locator={"sheet_name": sheet_name, "row_index": row_index},
                        record_id=_clean(mapping.get("id") or mapping.get("record_id") or f"{sheet_name}:{row_index}"),
                    )
                )
        workbook.close()
    except Exception:
        return ""
    return serialize_records(records)


def load_any_file(file_path: str) -> Document:
    extension = Path(file_path).suffix.lower()
    metadata = {
        "source_name": Path(file_path).name,
        "source_type": extension.lstrip("."),
        "load_error": None,
    }

    try:
        if extension == ".txt":
            content = load_txt(file_path)
        elif extension == ".pdf":
            content = load_pdf(file_path)
        elif extension == ".docx":
            content = load_docx(file_path)
        elif extension == ".doc":
            content = ""
            metadata["load_error"] = "Legacy .doc is unsupported; convert it to .docx first."
        elif extension == ".json":
            content = load_json(file_path)
        elif extension == ".csv":
            content = load_csv(file_path)
        elif extension == ".xlsx":
            content = load_xlsx(file_path)
        elif extension == ".xls":
            content = ""
            metadata["load_error"] = "Legacy .xls is unsupported; convert it to .xlsx first."
        elif extension in (".md", ".markdown"):
            content = load_txt(file_path)
        else:
            content = ""
            metadata["load_error"] = f"Unsupported file format: {extension or 'unknown'}"
    except Exception as exc:
        content = ""
        metadata["load_error"] = str(exc)

    if not content.strip() and not metadata.get("load_error"):
        metadata["load_error"] = "No extractable text found. The file may be scanned, empty, or malformed."

    return Document(content, metadata)
