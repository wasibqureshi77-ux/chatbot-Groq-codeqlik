import re
import os
import json
import csv
import io
import numpy as np
from rank_bm25 import BM25Okapi
from sklearn.metrics.pairwise import cosine_similarity

from database import knowledge_sources_collection, knowledge_chunks_collection
from rag.query_analyzer import analyze_query
from rag.ranker import rerank_chunks

# --- Third-party Library Safe Imports ---
try:
    import pypdf
except ImportError:
    pypdf = None

try:
    import docx
except ImportError:
    docx = None

try:
    import openpyxl
except ImportError:
    openpyxl = None

sources_collection = knowledge_sources_collection
chunks_collection = knowledge_chunks_collection

USE_METADATA_RAG = os.getenv("USE_METADATA_RAG", "false").lower() == "true"


# ==========================================
# 1. DOCUMENT CLASS & OPTIMIZED LOADERS
# ==========================================

class Document:
    def __init__(self, content: str, metadata: dict = None):
        self.content = content
        self.metadata = metadata or {}


def load_txt(file_path: str) -> str:
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def load_pdf(file_path: str) -> str:
    if pypdf:
        text = []
        try:
            reader = pypdf.PdfReader(file_path)
            for page_idx, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text:
                    text.append(f"=== PAGE_BREAK: {page_idx + 1} ===\n{page_text}")
            return "\n".join(text)
        except Exception as e:
            return f"[PDF Load Error: {str(e)}]"
    else:
        try:
            with open(file_path, "rb") as f:
                content = f.read()
            ascii_text = "".join(chr(b) if (32 <= b <= 126 or b in (10, 13)) else " " for b in content)
            return f"[pypdf not installed. Read raw bytes fallback]\n{ascii_text[:4000]}"
        except Exception:
            return "[PDF Read Error: pypdf not available]"


def load_docx(file_path: str) -> str:
    if docx:
        try:
            doc = docx.Document(file_path)
            content = []
            for para in doc.paragraphs:
                if para.text.strip():
                    content.append(para.text.strip())
            
            for table_idx, table in enumerate(doc.tables):
                content.append(f"--- Table {table_idx + 1} ---")
                headers = []
                for row_idx, row in enumerate(table.rows):
                    cells = [cell.text.strip() for cell in row.cells]
                    clean_cells = []
                    for c in cells:
                        if not clean_cells or c != clean_cells[-1]:
                            clean_cells.append(c)
                            
                    if row_idx == 0:
                        headers = clean_cells
                        content.append("Headers: " + " | ".join(headers))
                    else:
                        content.append("Row: " + " | ".join(
                            f"{headers[i] if i < len(headers) else f'Col{i}'}: {val}" 
                            for i, val in enumerate(clean_cells)
                        ))
            return "\n\n".join(content)
        except Exception as e:
            return f"[DOCX Load Error: {str(e)}]"
    else:
        return f"[python-docx not installed. Unable to load Word file {os.path.basename(file_path)}]"


def load_json(file_path: str) -> str:
    """
    RAG Highly Dense Context Optimization for Dynamic JSON layout.
    Injects context into every structural line for precise vector retrieval.
    """
    try:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            data = json.load(f)
        
        output = []
        
        def parse_element(element, context_summary="", prefix=""):
            if isinstance(element, dict):
                local_context = context_summary
                for identity_key in ['name', 'title', 'id', 'record_id', 'type']:
                    if identity_key in element and not isinstance(element[identity_key], (dict, list)):
                        local_context = f"{identity_key} '{element[identity_key]}'"
                        break
                
                for k, v in element.items():
                    current_key = f"{prefix}.{k}" if prefix else k
                    if isinstance(v, (dict, list)):
                        parse_element(v, local_context, current_key)
                    else:
                        if v is not None and str(v).strip():
                            ctx_str = f"Regarding {local_context}: " if local_context else ""
                            output.append(f"{ctx_str}The field '{current_key}' or '{k}' has the exact value: '{v}'")
            
            elif isinstance(element, list):
                for idx, item in enumerate(element):
                    current_prefix = f"{prefix}[{idx}]" if prefix else f"Record_{idx}"
                    parse_element(item, context_summary, current_prefix)
            else:
                if element is not None and str(element).strip():
                    output.append(f"Value for {prefix or 'element'}: '{element}'")

        parse_element(data)
        return "\n".join(output) if output else "[Empty JSON Content]"
    except Exception as e:
        return f"[JSON Load Error: {str(e)}]"


def load_csv(file_path: str) -> str:
    """
    RAG Highly Dense Context Optimization for dynamic CSV rows.
    Ensures ID and entity name mappings are repeated on every attribute line.
    """
    try:
        output = []
        filename = os.path.basename(file_path)
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            sample = f.read(2048)
            f.seek(0)
            try:
                dialect = csv.Sniffer().sniff(sample)
                reader = csv.reader(f, dialect)
            except Exception:
                reader = csv.reader(f)
                
            rows = list(reader)
            if not rows:
                return "[Empty CSV Content]"
            
            headers = [h.strip() if h.strip() else f"Column_{i+1}" for i, h in enumerate(rows[0])]
            
            for idx, row in enumerate(rows[1:]):
                if not any(cell.strip() for cell in row):
                    continue
                
                row_identity = f"Record Row {idx + 1}"
                for i, cell in enumerate(row):
                    if i < len(headers) and headers[i].lower() in ['name', 'title', 'id', 'record id', 'record_id']:
                        if cell.strip():
                            row_identity = f"Entity ({headers[i]}: {cell.strip()})"
                            break

                row_items = []
                for i, cell in enumerate(row):
                    val = cell.strip()
                    if val:
                        header_name = headers[i] if i < len(headers) else f"Column_{i+1}"
                        row_items.append(f"For {row_identity} in {filename} -> the '{header_name}' is '{val}'")
                
                if row_items:
                    output.extend(row_items)
                    
        return "\n".join(output)
    except Exception as e:
        return f"[CSV Load Error: {str(e)}]"


def load_xlsx(file_path: str) -> str:
    """
    RAG Highly Dense Context Optimization for Excel sheets.
    """
    if openpyxl:
        try:
            wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
            output = []
            filename = os.path.basename(file_path)
            for sheet in wb.sheetnames:
                ws = wb[sheet]
                rows = list(ws.iter_rows(values_only=True))
                if not rows:
                    continue
                
                headers = [str(cell).strip() if cell is not None and str(cell).strip() else f"Col_{i+1}" for i, cell in enumerate(rows[0])]
                
                for idx, row in enumerate(rows[1:]):
                    if any(cell is not None for cell in row):
                        row_identity = f"Row {idx+2}"
                        for i, cell in enumerate(row):
                            if i < len(headers) and headers[i].lower() in ['name', 'title', 'id', 'record id', 'record_id']:
                                if cell is not None and str(cell).strip():
                                    row_identity = f"Entity ({headers[i]}: {str(cell).strip()})"
                                    break
                                    
                        row_items = []
                        for i, cell in enumerate(row):
                            if cell is not None and str(cell).strip():
                                header_name = headers[i] if i < len(headers) else f"Col_{i+1}"
                                row_items.append(f"In Sheet '{sheet}' of '{filename}', for {row_identity} -> '{header_name}' is '{str(cell).strip()}'")
                        
                        if row_items:
                            output.extend(row_items)
            return "\n".join(output)
        except Exception as e:
            return f"[Excel Load Error: {str(e)}]"
    else:
        return f"[openpyxl not installed. Unable to load Excel file {os.path.basename(file_path)}]"


def load_any_file(file_path: str) -> Document:
    ext = os.path.splitext(file_path)[1].lower()
    metadata = {
        "source_name": os.path.basename(file_path),
        "source_type": ext.replace(".", "")
    }
    
    if ext == ".txt":
        content = load_txt(file_path)
    elif ext == ".pdf":
        content = load_pdf(file_path)
    elif ext in (".docx", ".doc"):
        content = load_docx(file_path)
    elif ext == ".json":
        content = load_json(file_path)
    elif ext == ".csv":
        content = load_csv(file_path)
    elif ext in (".xlsx", ".xls"):
        content = load_xlsx(file_path)
    elif ext in (".md", ".markdown"):
        content = load_txt(file_path)
    else:
        try:
            content = load_txt(file_path)
        except Exception:
            content = f"[Unsupported file format: {ext}]"

    return Document(content, metadata)


# ==========================================
# 2. INTENT & METADATA CONFIGURATION
# ==========================================

def detect_query_topic(query: str) -> (str, bool):
    q = query.lower()
    if any(w in q for w in ["technology", "technologies", "tech stack", "languages", "frameworks", "python", "javascript", "react", "node", "mongodb", "postgresql", "mysql", "fastapi", "langgraph", "groq", "llama"]):
        return "technologies", True
    if any(w in q for w in ["pricing", "price", "cost", "budget", "quote", "charges", "rate", "fees"]):
        return "pricing", True
    if any(w in q for w in ["contact", "address", "phone", "email", "office", "location", "reach us", "get in touch"]):
        return "contact", True
    if any(w in q for w in ["portfolio", "projects", "case study", "case studies", "our work", "clients", "track record", "delivered"]):
        return "portfolio", True
    if any(w in q for w in ["faq", "frequently asked questions", "common questions"]):
        return "faq", True
    if any(w in q for w in ["policy", "policies", "refund", "privacy", "terms of service", "cancellation", "tos"]):
        return "policies", True
    if any(w in q for w in ["service", "services", "provide", "offer", "specialty", "expertise", "custom development", "solution"]):
        return "services", True
    return "general", False


def detect_query_service(query: str) -> (str, bool):
    q = query.lower()
    if any(w in q for w in ["website", "web app", "web development", "react", "frontend", "html", "css", "static site"]):
        return "website", True
    if any(w in q for w in ["mobile app", "ios", "android", "flutter", "react native", "swift", "kotlin", "mobile application"]):
        return "mobile_app", True
    if any(w in q for w in ["ecommerce", "e-commerce", "shopify", "woocommerce", "online store", "shopping cart", "payment gateway"]):
        return "ecommerce", True
    if any(w in q for w in ["crm", "erp", "dashboard", "admin panel", "sales force", "hubspot", "management system"]):
        return "crm", True
    if any(w in q for w in ["ai automation", "chatbot", "rag", "langgraph", "agentic", "openai", "groq", "llama", "artificial intelligence", "automation script"]):
        return "ai_automation", True
    if any(w in q for w in ["software", "custom software", "saas", "cloud services", "backend", "database", "api"]):
        return "software", True
    return "general", False


def build_metadata_filter(detected_intent, detected_topic, topic_confident, detected_service, service_confident, filter_level=0):
    if filter_level >= 3:
        return {}
    conditions = []
    if filter_level < 3:
        conditions.append({
            "$or": [
                {"intent_scope": {"$in": [detected_intent, "all"]}},
                {"intent_scope": {"$exists": False}}
            ]
        })
    if filter_level < 2 and topic_confident:
        conditions.append({
            "$or": [
                {"topic": {"$in": [detected_topic, "general"]}},
                {"$and": [{"topic": {"$exists": False}}, {"category": {"$in": [detected_topic, "general"]}}]},
                {"$and": [{"topic": {"$exists": False}}, {"category": {"$exists": False}}]}
            ]
        })
    if filter_level < 1 and service_confident:
        conditions.append({
            "$or": [
                {"service": {"$in": [detected_service, "general"]}},
                {"service": {"$exists": False}}
            ]
        })
    if not conditions:
        return {}
    if len(conditions) == 1:
        return conditions[0]
    return {"$and": conditions}


def get_filtered_candidates(active_ids, identifiers, specific_keywords, generic_keywords, metadata_filter):
    candidates = []
    candidate_ids = set()
    
    def merge_filters(base, meta):
        if not meta:
            return base
        return {"$and": [base, meta]}
        
    target_pool_size = 80
        
    if identifiers:
        id_filter = {"source_id": {"$in": active_ids}}
        regex_conditions = []
        for kw in identifiers:
            regex_conditions.extend([
                {"chunk_text": re.compile(re.escape(kw), re.I)},
                {"title": re.compile(re.escape(kw), re.I)},
                {"category": re.compile(re.escape(kw), re.I)},
                {"keywords": re.compile(re.escape(kw), re.I)}
            ])
        id_filter["$or"] = regex_conditions
        query = merge_filters(id_filter, metadata_filter)
        
        id_candidates = list(chunks_collection.find(query).sort("created_at", -1).limit(30))
        for c in id_candidates:
            c_id = str(c["_id"])
            if c_id not in candidate_ids:
                candidates.append(c)
                candidate_ids.add(c_id)

    if len(candidates) < target_pool_size and specific_keywords:
        specific_filter = {"source_id": {"$in": active_ids}}
        regex_conditions = []
        for kw in specific_keywords:
            regex_conditions.extend([
                {"chunk_text": re.compile(re.escape(kw), re.I)},
                {"title": re.compile(re.escape(kw), re.I)},
                {"category": re.compile(re.escape(kw), re.I)},
                {"keywords": re.compile(re.escape(kw), re.I)}
            ])
        specific_filter["$or"] = regex_conditions
        query = merge_filters(specific_filter, metadata_filter)
        
        specific_candidates = list(chunks_collection.find(query).sort("created_at", -1).limit(50))
        for c in specific_candidates:
            c_id = str(c["_id"])
            if c_id not in candidate_ids:
                candidates.append(c)
                candidate_ids.add(c_id)

    if len(candidates) < target_pool_size and generic_keywords:
        generic_filter = {"source_id": {"$in": active_ids}}
        regex_conditions = []
        for kw in generic_keywords:
            regex_conditions.extend([
                {"chunk_text": re.compile(re.escape(kw), re.I)},
                {"title": re.compile(re.escape(kw), re.I)},
                {"category": re.compile(re.escape(kw), re.I)},
                {"keywords": re.compile(re.escape(kw), re.I)}
            ])
        generic_filter["$or"] = regex_conditions
        query = merge_filters(generic_filter, metadata_filter)
        
        generic_candidates = list(chunks_collection.find(query).sort("created_at", -1).limit(target_pool_size - len(candidates)))
        for c in generic_candidates:
            c_id = str(c["_id"])
            if c_id not in candidate_ids:
                candidates.append(c)
                candidate_ids.add(c_id)

    if not candidates:
        query = merge_filters({"source_id": {"$in": active_ids}}, metadata_filter)
        candidates = list(chunks_collection.find(query).sort("created_at", -1).limit(40))
        
    return candidates


def validate_context_relevance(query: str, chunks: list[dict], threshold: float = 0.18) -> bool:
    if not chunks:
        return False
    top_score = chunks[0].get("final_relevance_score", 0.0)
    return top_score >= threshold


# ==========================================
# 3. OPTIMIZED RETRIEVER LOGIC
# ==========================================

def retrieve_company_context_details(query: str, limit: int = 4, intent: str = None) -> dict:
    # 1. Query Analysis
    analysis = analyze_query(query)
    if not analysis.get("should_use_rag", True):
        return {
            "context_text": "No RAG context needed for general chat.",
            "confidence": 0.0,
            "sources": []
        }

    # 2. Get active sources
    active_sources = list(sources_collection.find({"enabled": True}, {"_id": 1, "title": 1}))
    if not active_sources:
        return {
            "context_text": "No active knowledge sources are configured. Respond with general polite support protocols.",
            "confidence": 0.0,
            "sources": []
        }

    active_ids = [str(s["_id"]) for s in active_sources]
    active_titles = {str(s["_id"]): s.get("title", "Unknown Source") for s in active_sources}

    # 3. Token & Identifier Extraction (Alphanumeric and short string handling)
    keywords = list(analysis.get("keywords", []))
    query_terms = re.findall(r'[a-zA-Z0-9\-]{2,}', query) # Length decreased to 2 to safely catch 'id'
    
    for term in query_terms:
        term_lower = term.lower()
        if term_lower not in [k.lower() for k in keywords] and term_lower not in {"give", "details", "order", "what", "who", "where", "how", "with", "from", "that", "this"}:
            keywords.append(term)

    identifiers = []
    for num in re.findall(r'\b\d{2,}\b', query):
        identifiers.append(num)
    for code in re.findall(r'\b[a-zA-Z0-9]+-\d+\b', query):
        identifiers.append(code)

    # REMOVED 'id' from GENERIC_WORDS so it stays as an important query token
    GENERIC_WORDS = {"give", "details", "order", "what", "who", "where", "how", "with", "from", "that", "this", "company", "service", "project", "client", "support", "hiring", "system", "website", "application", "business", "information"}
    
    specific_keywords = [k for k in keywords if k.lower() not in GENERIC_WORDS and k not in identifiers]
    generic_keywords = [k for k in keywords if k.lower() in GENERIC_WORDS and k not in identifiers]

    # 4. Metadata Guard Filter Loop (Protects against wrong manual inputs)
    candidates = []
    if USE_METADATA_RAG:
        detected_topic, topic_confident = detect_query_topic(query)
        detected_service, service_confident = detect_query_service(query)
        
        intent_map = {
            "client_lead": "client", "customer_support": "support", "hiring_support": "hiring",
            "general_chat": "greet", "client": "client", "support": "support",
            "hiring": "hiring", "greet": "greet", "all": "all"
        }
        detected_intent = intent_map.get(intent, "all")
        
        analysis["detected_query_topic"] = detected_topic
        analysis["detected_query_service"] = detected_service
        analysis["detected_intent_scope"] = detected_intent
        
        best_candidates = []
        for level in range(4):
            meta_filter = build_metadata_filter(
                detected_intent, detected_topic, topic_confident,
                detected_service, service_confident, filter_level=level
            )
            level_candidates = get_filtered_candidates(active_ids, identifiers, specific_keywords, generic_keywords, meta_filter)
            
            # Agar accurate candidates mil gaye toh loop break karo
            if len(level_candidates) >= 15:
                best_candidates = level_candidates
                break
            best_candidates = level_candidates

        # CRITICAL PROTECTION: Agar galat metadata ki wajah se chunks khali bache, bypass filters completely
        if len(best_candidates) < 5:
            best_candidates = get_filtered_candidates(active_ids, identifiers, specific_keywords, generic_keywords, {})
            
        candidates = best_candidates
    else:
        candidates = get_filtered_candidates(active_ids, identifiers, specific_keywords, generic_keywords, {})

    if not candidates:
        return {
            "context_text": "No highly relevant knowledge found.",
            "confidence": 0.0,
            "sources": []
        }
        
    chunks_text = [c.get("chunk_text", "") for c in candidates]

    # --- Hybrid Search Scoring Engine ---
    # BM25
    tokenized_chunks = [re.findall(r'[\w-]+', chunk.lower()) for chunk in chunks_text]
    tokenized_query = re.findall(r'[\w-]+', query.lower())
    bm25 = BM25Okapi(tokenized_chunks)
    bm25_scores = bm25.get_scores(tokenized_query)
    bm25_indices = np.argsort(bm25_scores)[::-1][:25].tolist()

    # Semantic Embedding Match
    from rag.embeddings import Embeddings
    embeddings_model = Embeddings()
    query_embedding = embeddings_model.embed_query(query)
    doc_embeddings = embeddings_model.embed_documents(chunks_text)
    similarities = cosine_similarity([query_embedding], doc_embeddings)[0]
    semantic_indices = similarities.argsort()[::-1][:30].tolist()

    # Entity & Internal Key Substring Matching
    entity_indices = []
    query_entities = analysis.get("entities", [])
    for idx, chunk in enumerate(chunks_text):
        chunk_lower = chunk.lower()
        for kw in specific_keywords:
            if kw.lower() in chunk_lower:
                entity_indices.append(idx)
                break
        for ent in query_entities:
            ent_text = ent["text"] if isinstance(ent, dict) else str(ent)
            if ent_text.lower() in chunk_lower and idx not in entity_indices:
                entity_indices.append(idx)
                break

    # Exact Match System
    exact_indices = []
    query_terms_clean = re.findall(r'[a-zA-Z0-9\-]{2,}', query.lower())
    for idx, chunk in enumerate(chunks_text):
        chunk_lower = chunk.lower()
        for term in query_terms_clean:
            if term in chunk_lower:
                exact_indices.append(idx)
                break

    # Score Weight Merger
    scores = {}
    for idx in semantic_indices:
        scores[idx] = scores.get(idx, 0.0) + (similarities[idx] * 0.5)

    for rank, idx in enumerate(bm25_indices):
        bm25_score = ((len(bm25_indices) - rank) / len(bm25_indices))
        scores[idx] = scores.get(idx, 0.0) + (bm25_score * 0.3)

    for idx in set(entity_indices):
        scores[idx] = scores.get(idx, 0.0) + 0.35  # Extra structural tag boost

    for idx in set(exact_indices):
        scores[idx] = scores.get(idx, 0.0) + 0.4

    ranked_indices = sorted(scores.items(), key=lambda x: x[1], reverse=True)
    
    ranked_candidates = []
    for idx, hybrid_score in ranked_indices[:20]:
        candidate = dict(candidates[idx])
        candidate["relevance_score"] = min(hybrid_score, 1.0)
        ranked_candidates.append(candidate)

    # 5. Reranking
    reranked = rerank_chunks(ranked_candidates, query, analysis)

    # 6. Context Validation
    is_valid = validate_context_relevance(query, reranked, threshold=0.18)
    if not is_valid:
        return {
            "context_text": "No highly relevant knowledge found.",
            "confidence": 0.0,
            "sources": []
        }

    top_chunks = reranked[:limit]
    context_blocks = []
    sources_used = []
    
    for chunk in top_chunks:
        src_id = chunk.get("source_id")
        source_name = active_titles.get(src_id, chunk.get("source_name", "Unknown Source"))
        source_type = chunk.get("source_type", "Text")
        category = chunk.get("category", "General")
        text = chunk.get("chunk_text", "").strip()
        
        meta_str = f"Source: {source_name} ({source_type}) | Category: {category}"
        meta = chunk.get("metadata", {})
        if "page_number" in meta:
            meta_str += f" | Page: {meta['page_number']}"
        if "sheet_name" in meta:
            meta_str += f" | Sheet: {meta['sheet_name']}"
            
        context_blocks.append(f"=== {meta_str} ===\n{text}")
        if source_name not in sources_used:
            sources_used.append(source_name)

    top_confidence = top_chunks[0].get("final_relevance_score", 0.0) if top_chunks else 0.0

    return {
        "context_text": "\n\n".join(context_blocks),
        "confidence": float(top_confidence),
        "sources": sources_used
    }


def retrieve_company_context(query: str, limit: int = 4) -> str:
    """
    Wrapper for backward compatibility.
    """
    details = retrieve_company_context_details(query, limit)
    return details["context_text"]